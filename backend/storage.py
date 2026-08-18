"""
Secure file storage service for clinical documents.
Files are stored locally with secure access control.
"""
import os
import re
import hashlib
import aiofiles
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple
from fastapi import UploadFile, HTTPException

# Storage configuration
STORAGE_ROOT = os.getenv("STORAGE_ROOT", os.path.join(os.path.dirname(__file__), "uploads"))
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

ALLOWED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".txt": "text/plain",
}

# Audio file extensions for Session Intelligence
AUDIO_EXTENSIONS = {
    ".mp3": ["audio/mpeg", "audio/mp3"],
    ".wav": ["audio/wav", "audio/x-wav", "audio/wave"],
    ".m4a": ["audio/m4a", "audio/x-m4a", "audio/mp4"],
    ".mp4": ["audio/mp4", "video/mp4"],
    # Browser MediaRecorder (used by the voice-profile "Record" flow) only emits webm/ogg
    ".webm": ["audio/webm", "video/webm"],
    ".ogg": ["audio/ogg", "audio/x-ogg"],
}


def ensure_storage_exists():
    """Create storage directory if it doesn't exist."""
    Path(STORAGE_ROOT).mkdir(parents=True, exist_ok=True)


def get_patient_storage_path(patient_id: str) -> Path:
    """Get the storage path for a patient's documents."""
    path = Path(STORAGE_ROOT) / patient_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def validate_file_type(filename: str, content_type: str) -> Tuple[bool, str]:
    """Validate file type based on extension and content type."""
    ext = Path(filename).suffix.lower()
    
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"File type '{ext}' is not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS.keys())}"
    
    expected_mime = ALLOWED_EXTENSIONS[ext]
    
    # Handle multiple valid MIME types for certain extensions
    valid_mimes = {
        ".jpg": ["image/jpeg", "image/jpg"],
        ".jpeg": ["image/jpeg", "image/jpg"],
    }
    
    if ext in valid_mimes:
        if content_type not in valid_mimes[ext]:
            return False, f"Invalid content type for {ext} file"
    elif content_type != expected_mime:
        # Allow generic octet-stream as fallback
        if content_type != "application/octet-stream":
            return False, f"Content type mismatch for {ext} file"
    
    return True, ""


async def compute_file_hash(file_content: bytes) -> str:
    """Compute SHA-256 hash of file content."""
    return hashlib.sha256(file_content).hexdigest()


async def save_file(
    file: UploadFile,
    patient_id: str,
    document_id: str,
) -> Tuple[str, int, str]:
    """
    Save uploaded file securely.
    
    Returns: (storage_path, file_size, file_hash)
    """
    ensure_storage_exists()
    
    # Read file content
    content = await file.read()
    file_size = len(content)
    
    # Validate file size
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    if file_size == 0:
        raise HTTPException(status_code=400, detail="Cannot upload empty file")
    
    # Compute hash for duplicate detection
    file_hash = await compute_file_hash(content)
    
    # Get file extension
    original_ext = Path(file.filename).suffix.lower()
    
    # Create storage path: uploads/{patient_id}/{document_id}{ext}
    patient_path = get_patient_storage_path(patient_id)
    storage_filename = f"{document_id}{original_ext}"
    storage_path = patient_path / storage_filename
    
    # Write file
    async with aiofiles.open(storage_path, "wb") as f:
        await f.write(content)
    
    # Return relative storage path
    relative_path = f"{patient_id}/{storage_filename}"
    
    return relative_path, file_size, file_hash


async def get_file_path(storage_path: str) -> Optional[Path]:
    """Get the absolute path to a stored file."""
    full_path = Path(STORAGE_ROOT) / storage_path
    if full_path.exists():
        return full_path
    return None


async def read_file(storage_path: str) -> Optional[bytes]:
    """Read file content from storage."""
    full_path = Path(STORAGE_ROOT) / storage_path
    if not full_path.exists():
        return None
    
    async with aiofiles.open(full_path, "rb") as f:
        return await f.read()


async def delete_file(storage_path: str) -> bool:
    """Delete a file from storage."""
    full_path = Path(STORAGE_ROOT) / storage_path
    if full_path.exists():
        full_path.unlink()
        return True
    return False


def get_mime_type(filename: str) -> str:
    """Get MIME type based on file extension."""
    ext = Path(filename).suffix.lower()
    return ALLOWED_EXTENSIONS.get(ext, "application/octet-stream")


# ═══════════════════════════════════════════════════════════════════════════════
#  AUDIO FILE STORAGE — Voice Profiles and Therapy Sessions
# ═══════════════════════════════════════════════════════════════════════════════

MAX_AUDIO_SIZE = 500 * 1024 * 1024  # 500MB for audio files


def validate_audio_file(filename: str, content_type: str) -> Tuple[bool, str]:
    """Validate audio file type based on extension and content type."""
    ext = Path(filename).suffix.lower()
    
    if ext not in AUDIO_EXTENSIONS:
        return False, f"Audio file type '{ext}' is not supported. Allowed types: {', '.join(AUDIO_EXTENSIONS.keys())}"
    
    expected_mimes = AUDIO_EXTENSIONS[ext]
    
    # Allow the file if content type matches any expected MIME type
    if content_type in expected_mimes:
        return True, ""
    
    # Also allow generic octet-stream as fallback
    if content_type == "application/octet-stream":
        return True, ""
    
    return False, f"Invalid content type for {ext} file. Expected one of: {expected_mimes}"


def get_voice_profile_storage_path(practitioner_id: str) -> Path:
    """Get the storage path for practitioner voice profiles."""
    path = Path(STORAGE_ROOT) / "voice_profiles" / practitioner_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_therapy_session_storage_path(patient_id: str) -> Path:
    """Get the storage path for patient therapy session recordings."""
    path = Path(STORAGE_ROOT) / "therapy_sessions" / patient_id
    path.mkdir(parents=True, exist_ok=True)
    return path


async def save_voice_profile(
    file: UploadFile,
    practitioner_id: str,
    profile_id: str,
) -> Tuple[str, int]:
    """
    Save voice profile audio file.
    
    Returns: (storage_path, file_size)
    """
    ensure_storage_exists()
    
    content = await file.read()
    file_size = len(content)
    
    if file_size > MAX_AUDIO_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Audio file size exceeds maximum allowed size of {MAX_AUDIO_SIZE // (1024*1024)}MB"
        )
    
    if file_size == 0:
        raise HTTPException(status_code=400, detail="Cannot upload empty file")
    
    original_ext = Path(file.filename).suffix.lower()
    profile_path = get_voice_profile_storage_path(practitioner_id)
    storage_filename = f"{profile_id}{original_ext}"
    storage_path = profile_path / storage_filename
    
    async with aiofiles.open(storage_path, "wb") as f:
        await f.write(content)
    
    relative_path = f"voice_profiles/{practitioner_id}/{storage_filename}"
    return relative_path, file_size


async def save_therapy_session_audio(
    file: UploadFile,
    patient_id: str,
    session_id: str,
) -> Tuple[str, int]:
    """
    Save therapy session audio file.
    
    Returns: (storage_path, file_size)
    """
    ensure_storage_exists()
    
    content = await file.read()
    file_size = len(content)
    
    if file_size > MAX_AUDIO_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Audio file size exceeds maximum allowed size of {MAX_AUDIO_SIZE // (1024*1024)}MB"
        )
    
    if file_size == 0:
        raise HTTPException(status_code=400, detail="Cannot upload empty file")
    
    original_ext = Path(file.filename).suffix.lower()
    session_path = get_therapy_session_storage_path(patient_id)
    storage_filename = f"{session_id}{original_ext}"
    storage_path = session_path / storage_filename
    
    async with aiofiles.open(storage_path, "wb") as f:
        await f.write(content)
    
    relative_path = f"therapy_sessions/{patient_id}/{storage_filename}"
    return relative_path, file_size


def get_audio_mime_type(filename: str) -> str:
    """Get MIME type for audio file based on extension."""
    ext = Path(filename).suffix.lower()
    mimes = AUDIO_EXTENSIONS.get(ext, [])
    return mimes[0] if mimes else "audio/mpeg"


# Session transcripts uploaded as a document rather than pasted as text —
# restricted to whatever _TEXT_EXTRACTORS (below) can actually pull text out
# of. No .doc/.xls here: legacy .doc isn't parsed (see extract_document_text).
TRANSCRIPT_DOCUMENT_EXTENSIONS = {
    ".pdf": ["application/pdf"],
    ".docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    ".txt": ["text/plain"],
}


def validate_transcript_document(filename: str, content_type: str) -> Tuple[bool, str]:
    """Validate a session-transcript file upload (PDF/DOCX/TXT only)."""
    ext = Path(filename).suffix.lower()

    if ext not in TRANSCRIPT_DOCUMENT_EXTENSIONS:
        allowed = ", ".join(TRANSCRIPT_DOCUMENT_EXTENSIONS.keys())
        return False, f"'{ext}' isn't supported for transcript upload. Allowed types: {allowed}"

    expected_mimes = TRANSCRIPT_DOCUMENT_EXTENSIONS[ext]
    if content_type in expected_mimes or content_type == "application/octet-stream":
        return True, ""

    return False, f"Invalid content type for {ext} file. Expected one of: {expected_mimes}"


async def save_therapy_session_document(
    file: UploadFile,
    patient_id: str,
    session_id: str,
) -> Tuple[str, int]:
    """
    Save a transcript document (PDF/DOCX/TXT) for a therapy session, alongside
    audio recordings in the same per-patient therapy_sessions folder.

    Returns: (storage_path, file_size)
    """
    ensure_storage_exists()

    content = await file.read()
    file_size = len(content)

    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE // (1024*1024)}MB"
        )

    if file_size == 0:
        raise HTTPException(status_code=400, detail="Cannot upload empty file")

    original_ext = Path(file.filename).suffix.lower()
    session_path = get_therapy_session_storage_path(patient_id)
    storage_filename = f"{session_id}{original_ext}"
    storage_path = session_path / storage_filename

    async with aiofiles.open(storage_path, "wb") as f:
        await f.write(content)

    relative_path = f"therapy_sessions/{patient_id}/{storage_filename}"
    return relative_path, file_size


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT TEXT EXTRACTION — feeds Clinical Intelligence (see clinical_intelligence.py)
# ═══════════════════════════════════════════════════════════════════════════════
#
# Best-effort extraction so uploaded reports (MMPI-2 today; ADHD/IQ/MCMI/etc. in the
# future, since those are administered externally and uploaded rather than built as
# in-app question flows) get their content analyzed, not just recorded as a file.

def _extract_pdf_text_sync(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        # Some PDFs (e.g. exports of justified/fully-spaced text) place each word
        # on its own content-stream line, so pypdf's line-based extraction returns
        # a newline after nearly every word instead of real paragraph breaks —
        # "we\nhad\na\nsession" rather than "we had a session". There's no reliable
        # signal in the extracted text to tell those apart from genuine line
        # breaks, so we collapse all in-page whitespace runs containing a newline
        # into a single space; page breaks (joined below) are preserved separately.
        page_text = re.sub(r"[ \t]*\n[ \t\n]*", " ", page_text).strip()
        if page_text:
            parts.append(page_text)
    return "\n\n".join(parts)


def _extract_docx_text_sync(path: Path) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    # Clinical/psych reports frequently put scores and tables in tables, not paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt_text_sync(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


_TEXT_EXTRACTORS = {
    ".pdf": _extract_pdf_text_sync,
    ".docx": _extract_docx_text_sync,
    ".txt": _extract_txt_text_sync,
}

# Formats a practitioner would reasonably expect to be readable "documents" but
# that extract_document_text can't parse (legacy binary .doc has no maintained
# pure-Python reader; .xls/.xlsx are spreadsheets, not prose). Used to tell
# these apart from images, which were never expected to yield text in the first
# place, so the UI can say *why* a document has no extracted text instead of
# just quietly extracting nothing.
UNSUPPORTED_TEXT_EXTENSIONS = {".doc", ".xls", ".xlsx"}


def is_unsupported_text_format(filename: str) -> bool:
    """True if `filename`'s extension is a known-unreadable-for-text format
    (see UNSUPPORTED_TEXT_EXTENSIONS) rather than just one that happened to
    yield no text (e.g. a scanned PDF)."""
    return Path(filename).suffix.lower() in UNSUPPORTED_TEXT_EXTENSIONS

# Matches a speaker turn label ("Client:", "Therapist:", etc.) — capitalized so it
# doesn't fire on the same words used mid-sentence in prose ("...with my previous
# therapist..."). Covers the labels this app's transcript documents use in practice.
_SPEAKER_LABEL_RE = re.compile(r"\s*\b(Client|Patient|Therapist|Counsel(?:l)?or|Doctor)\s*:\s*")


def _add_speaker_paragraph_breaks(text: str) -> str:
    """
    Two-party therapy transcripts are often exported as running prose with
    "Client:"/"Therapist:" turn labels inline rather than one turn per line —
    extracting them as-is (see _extract_pdf_text_sync) reads as a single wall of
    text. Insert a paragraph break before each recognized speaker label so the
    transcript displays one turn per paragraph, matching how these documents are
    actually laid out in their source (Word/PDF) form.
    """
    text = _SPEAKER_LABEL_RE.sub(lambda m: f"\n\n{m.group(1)}: ", text)
    return text.strip()


async def extract_document_text(storage_path: str) -> Optional[str]:
    """
    Extract plain text from a stored clinical document for Clinical Intelligence to
    analyze. Returns None for unsupported types (.doc, .xls/.xlsx, images — legacy
    .doc and spreadsheets aren't parsed, and images would need OCR) or when
    extraction finds nothing (e.g. a scanned/image-only PDF with no text layer).
    Callers should treat None as "no text available", not as an error — the document
    itself still uploads and is stored either way.
    """
    full_path = Path(STORAGE_ROOT) / storage_path
    if not full_path.exists():
        return None

    extractor = _TEXT_EXTRACTORS.get(full_path.suffix.lower())
    if not extractor:
        return None

    try:
        import anyio
        # pypdf/python-docx are sync and CPU-bound — keep them off the event loop
        text = await anyio.to_thread.run_sync(extractor, full_path)
    except Exception as e:
        print(f"Document text extraction error ({full_path.suffix}): {e}")
        return None

    text = (text or "").strip()
    if text:
        text = _add_speaker_paragraph_breaks(text)
    return text or None
